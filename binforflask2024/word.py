from docx import Document
from io import BytesIO
from datetime import datetime

# Function to generate a Word document dynamically
def generate_word_doc(email, company_name, website_url):
    doc_stream = BytesIO()
    document = Document()
    
    document.add_paragraph(f"Email: {email}")
    document.add_paragraph(f"Company Name: {company_name}")
    document.add_paragraph(f"Website URL: {website_url}")
    document.add_paragraph(f"Date Submitted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("\n" + "-" * 30 + "\n")
    
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream
