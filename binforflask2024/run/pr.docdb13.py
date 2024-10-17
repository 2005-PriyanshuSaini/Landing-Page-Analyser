from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from docx import Document
import os
from io import BytesIO
import uuid

# Initialize the Flask app
app = Flask(__name__)

# Default configuration for SQLite database (temporary URI)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///default.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # Disable unnecessary notifications
app.config['SECRET_KEY'] = 'your_secret_key'  # Required for session and flash messages

# Initialize SQLAlchemy
db = SQLAlchemy(app)

# Function to get the database URI based on the session
def get_db_uri():
    session_id = session.get('session_id')
    return f'sqlite:///data_{session_id}.db'

# Define the database model
class Data(db.Model):
    Sr = db.Column(db.Integer, primary_key=True)
    exampleInputEmail1 = db.Column(db.String(200), nullable=False)
    exampleInputName1 = db.Column(db.String(200), nullable=False)
    exampleInputUrl1 = db.Column(db.String(200), nullable=False)
    date_created = db.Column(db.DateTime, default=datetime.utcnow)

    def __repr__(self):
        return f'<Data {self.Sr}: {self.exampleInputEmail1}, {self.exampleInputName1}, {self.exampleInputUrl1}>'

# Ensure a new session ID is created and DB is initialized
@app.before_request
def ensure_session():
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())  # Generate a new unique session ID
    
    # Update database URI for the current session
    app.config['SQLALCHEMY_DATABASE_URI'] = get_db_uri()
    
    # Initialize the database schema for the current session
    with app.app_context():
        db.create_all()

# Function to generate a Word document dynamically
def generate_word_doc(email, company_name, website_url):
    doc_stream = BytesIO()
    document = Document()
    
    document.add_paragraph(f"Email: {email}")
    document.add_paragraph(f"Company Name: {company_name}")
    document.add_paragraph(f"Website URL: {website_url}")
    document.add_paragraph(f"Date Submitted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("\n" + "-" * 30 + "\n")
    
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Route for submitting data
@app.route('/', methods=['GET', 'POST'])
def submit_data():
    if request.method == 'POST':
        try:
            email = request.form['email']
            company_name = request.form['company_name']
            website_url = request.form['website_url']
            
            new_data = Data(exampleInputEmail1=email, exampleInputName1=company_name, exampleInputUrl1=website_url)
            db.session.add(new_data)
            db.session.commit()
            
            flash('Data successfully submitted!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error: {str(e)}', 'danger')

        return redirect(url_for('submit_data'))

    all_data = Data.query.all()
    return render_template('index.html', all_submissions=all_data)

# Route to allow downloading the Word document
@app.route('/download', methods=['POST'])
def download_doc():
    email = request.form['email']
    company_name = request.form['company_name']
    website_url = request.form['website_url']
    
    doc_stream = generate_word_doc(email, company_name, website_url)
    return send_file(doc_stream, as_attachment=True, download_name='submission.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)

