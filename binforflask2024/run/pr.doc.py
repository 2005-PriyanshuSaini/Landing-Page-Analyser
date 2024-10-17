from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from docx import Document
import os
from io import BytesIO

# Initialize the Flask app
app = Flask(__name__)

# Configuration for SQLite database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///data.db'  # Creates data.db in the project directory
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # Disable unnecessary notifications
app.config['SECRET_KEY'] = 'your_secret_key'  # Required for flash messages

# Initialize the SQLAlchemy object
db = SQLAlchemy(app)

# Define the database model
class Data(db.Model):
    Sr = db.Column(db.Integer, primary_key=True)  # Auto-incremented primary key
    exampleInputEmail1 = db.Column(db.String(200), nullable=False)  # Email input
    exampleInputName1 = db.Column(db.String(200), nullable=False)  # Company name input
    exampleInputUrl1 = db.Column(db.String(200), nullable=False)  # Website URL input
    date_created = db.Column(db.DateTime, default=datetime.utcnow)  # Auto-filled with the current date

    def __repr__(self):
        return f'<Data {self.Sr}: {self.exampleInputEmail1}, {self.exampleInputName1}, {self.exampleInputUrl1}>'

# Function to generate a Word document dynamically
def generate_word_doc(email, company_name, website_url):
    # Create an in-memory BytesIO object
    doc_stream = BytesIO()

    # Create a new Document object
    document = Document()
    
    # Add data to the document
    document.add_paragraph(f"Email: {email}")
    document.add_paragraph(f"Company Name: {company_name}")
    document.add_paragraph(f"Website URL: {website_url}")
    document.add_paragraph(f"Date Submitted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("\n" + "-" * 30 + "\n")  # Separator between entries
    
    # Save the document to the BytesIO stream
    document.save(doc_stream)
    
    # Move the stream position back to the start
    doc_stream.seek(0)
    
    return doc_stream

# Route for the main page
@app.route('/', methods=['GET', 'POST'])
def submit_data():
    if request.method == 'POST':
        # Capture form inputs
        try:
            email = request.form['email']
            company_name = request.form['company_name']
            website_url = request.form['website_url']
            
            # Create new entry for the database
            new_data = Data(exampleInputEmail1=email, exampleInputName1=company_name, exampleInputUrl1=website_url)
            db.session.add(new_data)
            db.session.commit()

            # Save the data to the database and flash a success message
            flash('Data successfully submitted!', 'success')
        except Exception as e:
            db.session.rollback()  # Undo changes on error
            flash(f'Error: {str(e)}', 'danger')  # Show an error message

        # Redirect to prevent form re-submission
        return redirect(url_for('submit_data'))

    # Retrieve all entries from the database
    all_data = Data.query.all()
    
    # Render the template and pass the retrieved data to the HTML
    return render_template('index.html', all_submissions=all_data)

# Route to allow downloading the Word document
@app.route('/download', methods=['POST'])
def download_doc():
    # Capture form inputs from the session
    email = request.form['email']
    company_name = request.form['company_name']
    website_url = request.form['website_url']
    
    # Generate the Word document
    doc_stream = generate_word_doc(email, company_name, website_url)
    
    # Send the document for download
    return send_file(doc_stream, as_attachment=True, download_name='submission.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# For testing purposes, additional route to display entries
@app.route('/show')
def show_entries():
    all_data = Data.query.all()  # Fetch all data from the database
    return render_template('show.html', all_submissions=all_data)

# Main entry point for running the Flask app
if __name__ == '__main__':
    # Ensure database tables are created before the app starts
    with app.app_context():
        db.create_all()  # Create database tables
    
    # Start the Flask application
    app.run(debug=True)
